from __future__ import annotations

import re
from pathlib import Path
from typing import Any, Dict, List


class ExportServiceError(RuntimeError):
    pass


def _safe_stem(text: str, *, max_len: int = 80) -> str:
    raw = str(text or '').strip()
    if not raw:
        return 'quiz'
    raw = re.sub(r"\s+", " ", raw)
    # Keep it mostly ASCII for cross-platform file names.
    raw = re.sub(r"[^A-Za-z0-9._-]+", "_", raw)
    raw = raw.strip("._-") or "quiz"
    return raw[:max_len]


def export_quiz_to_docx(
    quiz: Dict[str, Any],
    out_path: Path,
    *,
    answer_title: str = "Answer key",
    include_answer_text: bool = True,
    include_explanations: bool = False,
    include_images: bool = True,
) -> Path:
    """Export quiz dict (as returned by get_quiz_with_questions) into a .docx.

    Questions go first, answers are placed only in the final section.
    """

    try:
        from docx import Document
        from docx.shared import Inches
    except Exception as exc:  # pragma: no cover
        raise ExportServiceError("python-docx not installed. Install: pip install python-docx") from exc

    title = str(quiz.get("title") or "Quiz").strip() or "Quiz"
    quiz_id = int(quiz.get("id") or 0)

    questions = quiz.get("questions") or []
    if not isinstance(questions, list) or not questions:
        raise ExportServiceError("Quiz has no questions")

    doc = Document()

    doc.add_heading(title, level=1)

    meta_parts: List[str] = []
    if quiz_id:
        meta_parts.append(f"ID: {quiz_id}")
    meta_parts.append(f"Questions: {len(questions)}")
    if meta_parts:
        doc.add_paragraph(" | ".join(meta_parts))

    letters = "ABCD"

    for idx, q in enumerate(questions, start=1):
        q_text = str(q.get("question") or q.get("text") or "").strip()
        img_path = str(q.get("image_path") or "").strip()

        # Question line (bold)
        line = q_text
        if not line and (img_path or str(q.get("image_file_id") or "").strip()):
            line = "[Image question]"
        p = doc.add_paragraph()
        run = p.add_run(f"{idx}. {line}".strip())
        run.bold = True

        # Optional embedded image
        if include_images and img_path:
            ip = Path(img_path)
            if ip.exists() and ip.suffix.lower() in {".png", ".jpg", ".jpeg", ".webp"}:
                try:
                    # Keep image reasonably sized on page.
                    doc.add_picture(str(ip), width=Inches(5.8))
                except Exception:
                    pass

        opts = q.get("options") or []
        if not isinstance(opts, list):
            opts = []
        opts = [str(o) for o in opts][:4]

        # Always output 4 option lines (pad if needed)
        for i in range(4):
            opt = opts[i] if i < len(opts) else ""
            doc.add_paragraph(f"{letters[i]}) {opt}".rstrip())

        doc.add_paragraph("")

    # Answers section at the end
    doc.add_page_break()
    doc.add_heading(str(answer_title or "Answer key"), level=1)

    for idx, q in enumerate(questions, start=1):
        opts = q.get("options") or []
        if not isinstance(opts, list):
            opts = []
        opts = [str(o) for o in opts][:4]

        try:
            correct_index = int(q.get("correct_index") if q.get("correct_index") is not None else q.get("correct_answer") or 0)
        except Exception:
            correct_index = 0
        correct_index = max(0, min(3, int(correct_index)))

        letter = letters[correct_index]
        line = f"{idx}. {letter}"
        if include_answer_text and correct_index < len(opts):
            line += f") {opts[correct_index]}"
        doc.add_paragraph(line)

        if include_explanations:
            expl = str(q.get("explanation") or "").strip()
            if expl:
                doc.add_paragraph(expl)

    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    return out_path


def suggest_docx_filename(title: str, quiz_id: int) -> str:
    stem = _safe_stem(title, max_len=60)
    qid = int(quiz_id or 0)
    if qid:
        return f"quiz_{qid}_{stem}.docx"
    return f"quiz_{stem}.docx"
